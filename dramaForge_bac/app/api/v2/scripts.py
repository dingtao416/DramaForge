"""
DramaForge v2.0 — Scripts API
===============================
Endpoints for script generation, upload, editing, and approval.
"""

from __future__ import annotations

import asyncio as _asyncio
import json

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.database import AsyncSessionLocal as _AsyncSessionLocal, get_db
from app.models.project import Project, ProjectStep
from app.models.script import Script
from app.models.episode import Episode
from app.models.character import Character, CharacterRole
from app.models.scene import SceneLocation
from app.schemas.script import (
    ScriptGenerateRequest,
    ScriptDetail,
    ScriptUpdate,
    EpisodeUpdate,
)
from app.engines.script_engine import script_engine
from app.services.storage import storage
from app.services.user_model_resolver import user_model_resolver
from app.core.security import CurrentUser, DbSession
from app.core.billing_deps import require_credits

router = APIRouter()


def _sse_event(event: str, data: dict | str | None) -> str:
    """Format a Server-Sent Event string."""
    payload = data if isinstance(data, str) else json.dumps(data, ensure_ascii=False)
    return f"event: {event}\ndata: {payload}\n\n"


async def _get_project(project_id: int, db: AsyncSession) -> Project:
    project = await db.get(Project, project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


async def _get_script(project_id: int, db: AsyncSession) -> Script:
    stmt = (
        select(Script)
        .where(Script.project_id == project_id)
        .options(selectinload(Script.episodes))
    )
    result = await db.execute(stmt)
    script = result.scalar_one_or_none()
    if not script:
        raise HTTPException(status_code=404, detail="Script not found")
    return script


@router.post("/projects/{project_id}/script/generate", response_model=ScriptDetail)
async def generate_script(
    project_id: int,
    body: ScriptGenerateRequest,
    user: CurrentUser,
    db: DbSession,
):
    """AI-generate a structured script from user input."""
    # Consume credits for script generation
    await require_credits(db, user.id, "script_gen", description="剧本 AI 生成")
    await db.commit()  # release write lock before the AI call

    project = await _get_project(project_id, db)

    # Collect existing character & scene names (for dedup when preserving)
    existing_chars = await db.execute(
        select(Character.name).where(Character.project_id == project_id)
    )
    existing_char_names = set(existing_chars.scalars().all())
    existing_scenes = await db.execute(
        select(SceneLocation.name).where(SceneLocation.project_id == project_id)
    )
    existing_scene_names = set(existing_scenes.scalars().all())

    # Handle existing script
    existing = await db.execute(
        select(Script).where(Script.project_id == project_id).options(selectinload(Script.episodes))
    )
    old_script = existing.scalar_one_or_none()

    preserve = body.preserve_episodes

    if old_script and not preserve:
        await db.delete(old_script)
        await db.flush()
    elif old_script and preserve:
        # Update existing script in place — keep the row (and its episodes) alive
        pass  # We'll update fields after generation

    # Resolve user's configured chat model
    resolved = await user_model_resolver.resolve(db, user.id, "chat")

    # Generate via ScriptEngine
    result = await script_engine.create_from_text(
        user_input=body.user_input,
        project=project,
        genre=body.genre,
        total_episodes=body.total_episodes,
        duration=body.duration_per_episode,
        chat_model=resolved.model_id,
        chat_api_key=resolved.api_key,
        chat_base_url=resolved.base_url,
        chat_options=resolved.raw_params or {},
    )

    if old_script and preserve:
        # ── Preserve mode: update existing script, merge episodes ──
        script = old_script
        for key, value in result["script"].items():
            setattr(script, key, value)

        # Build a lookup of existing episodes by number
        old_episodes_by_num: dict[int, Episode] = {}
        for ep in old_script.episodes:
            old_episodes_by_num[ep.number] = ep

        new_episodes_data = result.get("episodes", [])
        kept_episode_ids = set()

        for ep_data in new_episodes_data:
            ep_num = ep_data.get("number", 0)
            if ep_num in old_episodes_by_num:
                # Update existing episode (preserves downstream segments/shots)
                existing_ep = old_episodes_by_num[ep_num]
                for key, value in ep_data.items():
                    setattr(existing_ep, key, value)
                kept_episode_ids.add(existing_ep.id)
            else:
                # Create new episode
                new_ep = Episode(script_id=script.id, **ep_data)
                db.add(new_ep)

        # Remove episodes that are no longer in the new script
        for ep_num, old_ep in old_episodes_by_num.items():
            if old_ep.id not in kept_episode_ids:
                await db.delete(old_ep)
    else:
        # ── Fresh mode: create new script ──
        script = Script(project_id=project_id, **result["script"])
        db.add(script)
        await db.flush()

        for ep_data in result["episodes"]:
            episode = Episode(script_id=script.id, **ep_data)
            db.add(episode)

    # Create Characters (skip if name already exists — preserve existing images)
    for ch_data in result["characters"]:
        if ch_data.get("name") in existing_char_names:
            continue  # Keep existing character with its generated images
        role_str = ch_data.pop("role", "supporting")
        try:
            role = CharacterRole(role_str)
        except ValueError:
            role = CharacterRole.SUPPORTING
        character = Character(
            project_id=project_id,
            role=role,
            **ch_data,
        )
        db.add(character)

    # Create Scenes (skip if name already exists — preserve existing images)
    for sc_data in result["scenes"]:
        if sc_data.get("name") in existing_scene_names:
            continue
        scene = SceneLocation(project_id=project_id, **sc_data)
        db.add(scene)

    await db.flush()
    await db.refresh(script, attribute_names=["episodes"])
    return script


# ── Background Generation Registry ──────────────────────────────────
# Maps project_id → {queue, task, status, content_preview}
# Survives SSE disconnects so generation continues even if user navigates away.

_gen_registry: dict[int, dict] = {}


async def _save_script_result(project_id: int, result: dict, preserve: bool = False) -> int:
    """Save generated script to DB using a fresh session. Returns script_id."""
    async with _AsyncSessionLocal() as sess:
        if preserve:
            # Update existing script in place
            stmt = select(Script).where(Script.project_id == project_id).options(selectinload(Script.episodes))
            r = await sess.execute(stmt)
            script = r.scalar_one_or_none()
            if script:
                for key, value in result["script"].items():
                    setattr(script, key, value)
                # Merge episodes: update matching, add new, remove old
                old_eps_by_num = {ep.number: ep for ep in script.episodes}
                kept_ids = set()
                for ep_data in result.get("episodes", []):
                    ep_num = ep_data.get("number", 0)
                    if ep_num in old_eps_by_num:
                        existing_ep = old_eps_by_num[ep_num]
                        for k, v in ep_data.items():
                            setattr(existing_ep, k, v)
                        kept_ids.add(existing_ep.id)
                    else:
                        new_ep = Episode(script_id=script.id, **ep_data)
                        sess.add(new_ep)
                for ep_num, old_ep in old_eps_by_num.items():
                    if old_ep.id not in kept_ids:
                        await sess.delete(old_ep)
                await sess.flush()
            else:
                # No existing script, create new
                script = Script(project_id=project_id, **result["script"])
                sess.add(script)
                await sess.flush()
                for ep_data in result["episodes"]:
                    episode = Episode(script_id=script.id, **ep_data)
                    sess.add(episode)
        else:
            script = Script(project_id=project_id, **result["script"])
            sess.add(script)
            await sess.flush()

            for ep_data in result["episodes"]:
                episode = Episode(script_id=script.id, **ep_data)
                sess.add(episode)

        # Characters: skip if name already exists
        existing_chars = await sess.execute(
            select(Character.name).where(Character.project_id == project_id)
        )
        existing_char_names = set(existing_chars.scalars().all())
        for ch_data in result["characters"]:
            if ch_data.get("name") in existing_char_names:
                continue
            role_str = ch_data.pop("role", "supporting")
            try:
                role = CharacterRole(role_str)
            except ValueError:
                role = CharacterRole.SUPPORTING
            character = Character(project_id=project_id, role=role, **ch_data)
            sess.add(character)

        # Scenes: skip if name already exists
        existing_scenes = await sess.execute(
            select(SceneLocation.name).where(SceneLocation.project_id == project_id)
        )
        existing_scene_names = set(existing_scenes.scalars().all())
        for sc_data in result["scenes"]:
            if sc_data.get("name") in existing_scene_names:
                continue
            scene = SceneLocation(project_id=project_id, **sc_data)
            sess.add(scene)

        await sess.commit()
        await sess.refresh(script, attribute_names=["episodes"])
        return script.id


@router.get("/projects/{project_id}/script/generate-status")
async def get_generate_status(project_id: int, user: CurrentUser):
    """
    Check if script generation is in progress for a project.
    Returns: {status: "idle"|"generating"|"complete"|"error", content_length?: int, ...}
    """
    entry = _gen_registry.get(project_id)
    if not entry:
        return {"status": "idle"}
    return {
        "status": entry["status"],
        "content_length": len(entry.get("content_preview", "")),
    }


@router.post("/projects/{project_id}/script/cancel-generation")
async def cancel_generation(project_id: int, user: CurrentUser):
    """
    Cancel an ongoing script generation for a project.
    The background task will stop and NOT save partial results to DB.
    """
    entry = _gen_registry.get(project_id)
    if entry and entry["status"] == "generating":
        entry["cancelled"] = True
        entry["status"] = "cancelled"
        return {"message": "Generation cancelled", "status": "cancelled"}
    return {"message": "No active generation to cancel", "status": entry["status"] if entry else "idle"}


@router.post("/projects/{project_id}/script/generate-stream")
async def generate_script_stream(
    project_id: int,
    body: ScriptGenerateRequest,
    user: CurrentUser,
    db: DbSession,
):
    """
    AI-generate a structured script with SSE streaming progress.

    Generation runs as a background task — it survives page navigation.
    Reconnect by calling this endpoint again or check status at /generate-status.

    SSE events:
        - delta      → {content: "chunk..."}
        - done       → {script_id, episode_count, character_count, scene_count}
        - error      → {message: "..."}
        - heartbeat  → {} (every 15s to keep connection alive)
    """
    await require_credits(db, user.id, "script_gen", description="剧本 AI 生成(流式)")

    project = await _get_project(project_id, db)

    # Resolve user's configured chat model
    resolved = await user_model_resolver.resolve(db, user.id, "chat")

    # ── Check for existing generation ──
    existing_entry = _gen_registry.get(project_id)
    if existing_entry and existing_entry["status"] == "generating":
        # Reconnect: replay accumulated content then continue streaming
        queue = existing_entry["queue"]
        accumulated = existing_entry.get("content_preview", "")

        async def replay_generator():
            # Send accumulated content as a single delta first
            if accumulated:
                yield _sse_event("delta", {"content": accumulated})
            # Then continue reading from the live queue
            try:
                while True:
                    try:
                        event = await _asyncio.wait_for(queue.get(), timeout=15.0)
                    except _asyncio.TimeoutError:
                        yield _sse_event("heartbeat", {})
                        continue
                    if event["type"] == "content":
                        yield _sse_event("delta", {"content": event["data"]})
                    elif event["type"] == "done":
                        yield _sse_event("done", {
                            "script_id": event.get("script_id"),
                            "episode_count": event.get("episode_count", 0),
                            "character_count": event.get("character_count", 0),
                            "scene_count": event.get("scene_count", 0),
                        })
                        break
                    elif event["type"] == "error":
                        yield _sse_event("error", {"message": event["data"]})
                        break
            except _asyncio.CancelledError:
                pass

        return StreamingResponse(
            replay_generator(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",
            },
        )

    # ── Start a new background generation ──
    preserve = body.preserve_episodes
    if not preserve:
        existing = await db.execute(
            select(Script).where(Script.project_id == project_id)
        )
        old_script = existing.scalar_one_or_none()
        if old_script:
            await db.delete(old_script)
            await db.flush()
    # Commit now so background task's DB session won't be blocked by write lock
    await db.commit()

    queue: _asyncio.Queue = _asyncio.Queue()
    entry = {
        "queue": queue,
        "status": "generating",
        "content_preview": "",
        "task": None,
        "result": None,
        "cancelled": False,
    }
    _gen_registry[project_id] = entry

    async def background_generator():
        """Runs in background — survives SSE disconnect, but checks for cancellation."""
        content_buf = ""
        try:
            async for event in script_engine.create_from_text_stream(
                user_input=body.user_input,
                project=project,
                genre=body.genre,
                total_episodes=body.total_episodes,
                duration=body.duration_per_episode,
                chat_model=resolved.model_id,
                chat_api_key=resolved.api_key,
                chat_base_url=resolved.base_url,
                chat_options=resolved.raw_params or {},
            ):
                # Check cancellation before processing each event
                if entry.get("cancelled"):
                    entry["status"] = "cancelled"
                    return
                if event["type"] == "content":
                    content_buf += event["data"]
                    entry["content_preview"] = content_buf
                    await queue.put(event)
                elif event["type"] == "error":
                    await queue.put(event)
                    entry["status"] = "error"
                    return
                elif event["type"] == "done":
                    # Do NOT save if generation was cancelled
                    if entry.get("cancelled"):
                        entry["status"] = "cancelled"
                        return
                    # Save to DB with fresh session
                    try:
                        script_id = await _save_script_result(project_id, event["data"], preserve)
                        event["data"]["script_id"] = script_id
                        event["data"]["episode_count"] = len(event["data"].get("episodes", []))
                        event["data"]["character_count"] = len(event["data"].get("characters", []))
                        event["data"]["scene_count"] = len(event["data"].get("scenes", []))
                    except Exception as exc:
                        await queue.put({"type": "error", "data": f"DB save failed: {exc}"})
                        entry["status"] = "error"
                        return
                    entry["result"] = event["data"]
                    await queue.put(event)
                    entry["status"] = "complete"
                    return
        except Exception as exc:
            await queue.put({"type": "error", "data": str(exc)})
            entry["status"] = "error"

    task = _asyncio.create_task(background_generator())
    entry["task"] = task

    async def event_generator():
        try:
            while True:
                try:
                    event = await _asyncio.wait_for(queue.get(), timeout=15.0)
                except _asyncio.TimeoutError:
                    yield _sse_event("heartbeat", {})
                    continue

                if event["type"] == "content":
                    yield _sse_event("delta", {"content": event["data"]})
                elif event["type"] == "error":
                    yield _sse_event("error", {"message": event["data"]})
                    break
                elif event["type"] == "done":
                    yield _sse_event("done", {
                        "script_id": event.get("script_id"),
                        "episode_count": event.get("episode_count", 0),
                        "character_count": event.get("character_count", 0),
                        "scene_count": event.get("scene_count", 0),
                    })
                    break
        except _asyncio.CancelledError:
            # Client disconnected — background task keeps running
            pass

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


from pydantic import BaseModel, Field


class ScriptParseResponse(BaseModel):
    filename: str
    file_type: str
    char_count: int
    full_text: str
    preview: str = Field("", description="First 500 chars for display")


@router.post("/scripts/parse", response_model=ScriptParseResponse)
async def parse_script_file(
    file: UploadFile = File(...),
):
    """
    Parse an uploaded script file (.docx / .doc / .txt) and return its text content.

    This is a preview-only endpoint — no project is created.
    Use this to let users review the parsed content before creating a project.
    """
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ("docx", "doc", "txt"):
        raise HTTPException(status_code=400, detail="仅支持 .docx / .doc / .txt 格式")

    # Save to temp file
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    try:
        full_text = await script_engine.parse_uploaded_file(tmp_path)
    finally:
        import os
        os.unlink(tmp_path)

    return ScriptParseResponse(
        filename=file.filename,
        file_type=ext,
        char_count=len(full_text),
        full_text=full_text,
        preview=full_text[:500],
    )


@router.post("/projects/{project_id}/script/upload", response_model=ScriptDetail)
async def upload_script(
    project_id: int,
    file: UploadFile = File(...),
    total_episodes: int = 1,
    db: AsyncSession = Depends(get_db),
):
    """Upload a .docx/.doc/.txt script file."""
    project = await _get_project(project_id, db)

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ("docx", "doc", "txt"):
        raise HTTPException(status_code=400, detail="仅支持 .docx / .doc / .txt 格式")

    # Save uploaded file
    dest = storage.upload_path(project_id, file.filename)
    content = await file.read()
    await storage.save_from_bytes(content, dest)

    # Parse via ScriptEngine
    result = await script_engine.create_from_docx(
        file_path=dest,
        project=project,
        total_episodes=total_episodes,
    )

    # Remove existing script
    existing = await db.execute(
        select(Script).where(Script.project_id == project_id)
    )
    old_script = existing.scalar_one_or_none()
    if old_script:
        await db.delete(old_script)
        await db.flush()

    # Create Script + Episodes
    script = Script(project_id=project_id, **result["script"])
    db.add(script)
    await db.flush()

    for ep_data in result["episodes"]:
        db.add(Episode(script_id=script.id, **ep_data))

    for ch_data in result["characters"]:
        role_str = ch_data.pop("role", "supporting")
        try:
            role = CharacterRole(role_str)
        except ValueError:
            role = CharacterRole.SUPPORTING
        db.add(Character(project_id=project_id, role=role, **ch_data))

    for sc_data in result["scenes"]:
        db.add(SceneLocation(project_id=project_id, **sc_data))

    await db.flush()
    await db.refresh(script, attribute_names=["episodes"])
    return script


@router.get("/projects/{project_id}/script", response_model=ScriptDetail)
async def get_script(
    project_id: int,
    db: AsyncSession = Depends(get_db),
):
    """Get the script for a project."""
    return await _get_script(project_id, db)


@router.put("/projects/{project_id}/script", response_model=ScriptDetail)
async def update_script(
    project_id: int,
    body: ScriptUpdate,
    db: AsyncSession = Depends(get_db),
):
    """Edit script metadata."""
    script = await _get_script(project_id, db)

    update_data = body.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        setattr(script, key, value)

    await db.flush()
    await db.refresh(script, attribute_names=["episodes"])
    return script


@router.put("/projects/{project_id}/episodes/{episode_id}", response_model=ScriptDetail)
async def update_episode(
    project_id: int,
    episode_id: int,
    body: EpisodeUpdate,
    db: AsyncSession = Depends(get_db),
):
    """Edit an episode's content."""
    episode = await db.get(Episode, episode_id)
    if not episode:
        raise HTTPException(status_code=404, detail="Episode not found")

    update_data = body.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        setattr(episode, key, value)

    await db.flush()
    return await _get_script(project_id, db)


@router.post("/projects/{project_id}/script/rewrite-narration", response_model=ScriptDetail)
async def rewrite_narration(
    project_id: int,
    db: AsyncSession = Depends(get_db),
):
    """Rewrite script from dialogue to narration style."""
    script = await _get_script(project_id, db)

    # Collect all episode content
    all_content = "\n\n".join(
        ep.content for ep in script.episodes if ep.content
    )
    if not all_content:
        raise HTTPException(status_code=400, detail="No script content to rewrite")

    narration = await script_engine.rewrite_to_narration(all_content)

    # Update the first episode's content (or all if single-episode)
    if len(script.episodes) == 1:
        script.episodes[0].content = narration
    else:
        # For multi-episode, rewrite each separately (simplified: store in raw)
        script.raw_content = narration

    await db.flush()
    await db.refresh(script, attribute_names=["episodes"])
    return script


@router.post("/projects/{project_id}/script/rewrite-narration-stream")
async def rewrite_narration_stream(
    project_id: int,
    db: DbSession,
    user: CurrentUser,
):
    """
    Rewrite script from dialogue to narration style with SSE streaming.

    SSE events:
        - delta   → {content: "chunk..."}
        - done    → {content: "full narration text"}
        - error   → {message: "..."}
    """
    script = await _get_script(project_id, db)

    # Collect all episode content
    all_content = "\n\n".join(
        ep.content for ep in script.episodes if ep.content
    )
    if not all_content:
        raise HTTPException(status_code=400, detail="No script content to rewrite")

    async def event_generator():
        try:
            async for event in script_engine.rewrite_to_narration_stream(all_content):
                event_type = event["type"]
                event_data = event["data"]

                if event_type == "content":
                    yield _sse_event("delta", {"content": event_data})
                elif event_type == "error":
                    yield _sse_event("error", {"message": event_data})
                    return
                elif event_type == "done":
                    narration = event_data

                    # Update script
                    if len(script.episodes) == 1:
                        script.episodes[0].content = narration
                    else:
                        script.raw_content = narration

                    await db.commit()

                    yield _sse_event("done", {"content": narration})

        except Exception as exc:
            yield _sse_event("error", {"message": str(exc)})

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/projects/{project_id}/script/approve", response_model=ScriptDetail)
async def approve_script(
    project_id: int,
    db: AsyncSession = Depends(get_db),
):
    """Approve the script and advance project to Step 2 (Assets)."""
    script = await _get_script(project_id, db)

    script.is_approved = True
    for ep in script.episodes:
        ep.is_approved = True

    # Advance project status
    project = await _get_project(project_id, db)
    project.status = ProjectStep.ASSETS

    await db.flush()
    await db.refresh(script, attribute_names=["episodes"])
    return script


# ═══════════════════════════════════════════════════════════════
# P2-2: Script export
# ═══════════════════════════════════════════════════════════════

@router.get("/projects/{project_id}/script/export")
async def export_script(
    project_id: int,
    format: str = Query("docx", description="Export format: docx or txt"),
    user: CurrentUser = None,
    db: AsyncSession = Depends(get_db),
):
    """Export the script as a downloadable file (DOCX or TXT)."""
    from fastapi.responses import Response
    from io import BytesIO

    script = await _get_script(project_id, db)
    project = await _get_project(project_id, db)

    if format == "txt":
        lines = [f"《{project.title}》剧本", "=" * 40, ""]
        lines.append(f"主角：{script.protagonist or '未指定'}")
        lines.append(f"类型：{script.genre or '未指定'}")
        lines.append(f"梗概：{script.synopsis or '未指定'}")
        lines.append(f"背景：{script.background or '未指定'}")
        lines.append("")

        for ep in script.episodes:
            lines.append(f"第{ep.number}集：{ep.title or ''}")
            lines.append("-" * 30)
            lines.append(ep.content or "（暂无内容）")
            lines.append("")

        content = "\n".join(lines)
        safe_title = "".join(c for c in project.title if c.isalnum() or c in "._- ")
        filename = f"{safe_title}.txt"

        return Response(
            content=content.encode("utf-8-sig"),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # DOCX export
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx is required for DOCX export")

    doc = Document()

    # Title
    title_para = doc.add_heading(f"《{project.title}》", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta table
    meta_table = doc.add_table(rows=4, cols=2)
    for i, (label, value) in enumerate([
        ("主角", script.protagonist or ""),
        ("类型", script.genre or ""),
        ("梗概", script.synopsis or ""),
        ("背景", script.background or ""),
    ]):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value

    doc.add_paragraph("")

    # Episodes
    for ep in script.episodes:
        doc.add_heading(f"第{ep.number}集：{ep.title or ''}", level=2)
        if ep.content:
            for line in ep.content.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("（暂无内容）")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = "".join(c for c in project.title if c.isalnum() or c in "._- ")
    filename = f"{safe_title}.docx"

    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
